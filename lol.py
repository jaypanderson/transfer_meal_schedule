import docx
from tkinter import filedialog
from transfer_meal_schedule import new_file_path


def copy_table_style(src_cell, dst_cell):
    # Copy cell dimensions
    dst_cell.width = src_cell.width
    dst_cell.height = src_cell.height

    # Copy cell text and paragraph formatting
    if src_cell.text:
        dst_cell.text = src_cell.text
        src_paragraph = src_cell.paragraphs[0]
        dst_paragraph = dst_cell.paragraphs[0]

        # Copy font style
        if src_paragraph.runs:
            src_run = src_paragraph.runs[0]
            dst_run = dst_paragraph.runs[0] if dst_paragraph.runs else dst_paragraph.add_run()

            dst_run.font.name = src_run.font.name
            dst_run.font.size = src_run.font.size
            dst_run.font.bold = src_run.font.bold
            dst_run.font.italic = src_run.font.italic
            dst_run.font.underline = src_run.font.underline
            dst_run.font.color.rgb = src_run.font.color.rgb

        # Copy paragraph alignment
        dst_paragraph.alignment = src_paragraph.alignment

    # Copy cell border styles (example for top border)
    dst_cell.top_border = src_cell.top_border
    # Repeat for other borders (left, right, bottom)

    # Copy cell shading (background color)
    dst_cell.shading.fill = src_cell.shading.fill

    # Cell margins (example for top margin)
    dst_cell.top_margin = src_cell.top_margin
    # Repeat for other margins (left, right, bottom)


path = filedialog.askopenfilename(title='原本を選択してください。')
original = docx.Document(path)
doc = docx.Document()
src_table = original.tables[0]
result_table = doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
for i, row in enumerate(src_table.rows):
    for j, cell in enumerate(row.cells):
        copy_table_style(cell, result_table.cell(i,j))

doc.save(new_file_path(path, added_text='test_test_test'))